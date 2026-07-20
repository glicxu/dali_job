from __future__ import annotations

from io import BytesIO
from typing import Literal

from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.modules.materials.schemas import CoverLetterOutput, TailoredResumeOutput

RenderFormat = Literal["pdf", "docx"]


def material_lines(material_type: str, content_data: dict) -> list[tuple[str, str]]:
    lines: list[tuple[str, str]] = []
    if material_type == "tailored_resume":
        content = TailoredResumeOutput.model_validate(content_data)
        if content.headline:
            lines.append(("title", content.headline.text))
        for label, field in (
            ("Summary", content.summary),
            ("Skills", content.skills),
            ("Experience", content.experience),
            ("Education", content.education),
            ("Certifications", content.certifications),
            ("Projects", content.projects),
        ):
            if not field:
                continue
            lines.append(("heading", label))
            lines.extend(("bullet", item.text) for item in field)
        return lines

    content = CoverLetterOutput.model_validate(content_data)
    lines.append(("paragraph", content.salutation))
    lines.extend(("paragraph", paragraph.text) for paragraph in content.paragraphs)
    lines.append(("paragraph", content.closing))
    return lines


def plain_text(material_type: str, content_data: dict) -> str:
    output: list[str] = []
    for kind, text in material_lines(material_type, content_data):
        output.append(f"- {text}" if kind == "bullet" else text)
    return "\n\n".join(output).strip()


def render_material(material_type: str, content_data: dict, output_format: RenderFormat) -> bytes:
    lines = material_lines(material_type, content_data)
    if output_format == "docx":
        return _render_docx(lines)
    if output_format == "pdf":
        return _render_pdf(lines)
    raise ValueError("Unsupported render format.")


def _render_docx(lines: list[tuple[str, str]]) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    for kind, text in lines:
        if kind == "title":
            document.add_heading(text, level=0)
        elif kind == "heading":
            document.add_heading(text, level=1)
        elif kind == "bullet":
            document.add_paragraph(text, style="List Bullet")
        else:
            document.add_paragraph(text)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _render_pdf(lines: list[tuple[str, str]]) -> bytes:
    output = BytesIO()
    styles = getSampleStyleSheet()
    story = []
    for kind, text in lines:
        style = styles["Title"] if kind == "title" else styles["Heading2"] if kind == "heading" else styles["BodyText"]
        rendered = f"• {text}" if kind == "bullet" else text
        story.append(Paragraph(rendered.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), style))
        story.append(Spacer(1, 0.08 * inch))
    SimpleDocTemplate(
        output,
        pagesize=LETTER,
        rightMargin=0.7 * inch,
        leftMargin=0.7 * inch,
        topMargin=0.7 * inch,
        bottomMargin=0.7 * inch,
        title="DaliJob application material",
    ).build(story)
    return output.getvalue()
